"""文档文件读取 —— 支持 PDF、DOCX、TXT 格式。

读取原始文本和元数据（页码、标题等），不负责清洗和切块。
"""

import hashlib
import logging
import os
from dataclasses import dataclass, field

from contracts.errors import (
    DOCUMENT_PARSE_FAILED,
    INVALID_FILE_TYPE,
    ServiceError,
)

logger = logging.getLogger("rag.document_loader")


@dataclass
class PageInfo:
    """单页信息。"""
    page_number: int
    text: str


@dataclass
class DocumentLoadResult:
    """文档加载结果（原始内容 + 元数据）。"""
    content: str
    pages: list[PageInfo] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    """读取 PDF、DOCX、TXT 文档，返回原始文本和页级元数据。"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    def load(self, file_path: str) -> DocumentLoadResult:
        """根据扩展名分派到对应的加载器。

        Raises:
            ServiceError(INVALID_FILE_TYPE): 不支持的文件格式。
            ServiceError(DOCUMENT_PARSE_FAILED): 文档解析失败。
        """
        if not os.path.isfile(file_path):
            raise ServiceError(
                code=DOCUMENT_PARSE_FAILED,
                message=f"文件不存在: {file_path}",
            )

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ServiceError(
                code=INVALID_FILE_TYPE,
                message=f"不支持的文件格式: {ext}，支持的格式: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}",
            )

        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        base_meta = {
            "filename": filename,
            "file_path": file_path,
            "file_size": file_size,
            "extension": ext,
        }

        try:
            if ext == ".pdf":
                result = self._load_pdf(file_path)
            elif ext == ".docx":
                result = self._load_docx(file_path)
            elif ext in {".txt", ".md"}:
                result = self._load_txt(file_path)
            else:
                # 已经在前面的检查中排除，这里是防御性代码
                raise ServiceError(code=INVALID_FILE_TYPE, message=f"不支持的文件格式: {ext}")
        except ServiceError:
            raise
        except Exception as e:
            logger.exception("Document parsing failed: %s", file_path)
            raise ServiceError(
                code=DOCUMENT_PARSE_FAILED,
                message=f"文档解析失败 ({filename}): {e}",
                details={"file_path": file_path, "error": str(e)},
            ) from e

        result.metadata.update(base_meta)
        logger.info("Loaded document: %s, chars=%d, pages=%d", filename, len(result.content), len(result.pages))
        return result

    # ── PDF ─────────────────────────────────────────────────

    def _load_pdf(self, file_path: str) -> DocumentLoadResult:
        """使用 PyMuPDF (fitz) 读取 PDF，fallback 到 pdfplumber / PyPDF2。"""
        pages: list[PageInfo] = []
        full_text_parts: list[str] = []

        # 优先使用 PyMuPDF
        try:
            import fitz  # type: ignore[import-untyped]
            doc = fitz.open(file_path)
            for i, page in enumerate(doc, start=1):
                text = page.get_text()
                pages.append(PageInfo(page_number=i, text=text))
                full_text_parts.append(text)
            doc.close()
            return DocumentLoadResult(
                content="\n\n".join(full_text_parts),
                pages=pages,
            )
        except ImportError:
            logger.debug("PyMuPDF not available, trying pdfplumber")
        except Exception as e:
            logger.warning("PyMuPDF failed: %s, trying fallback", e)

        # Fallback 1: pdfplumber
        try:
            import pdfplumber  # type: ignore[import-untyped]
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    pages.append(PageInfo(page_number=i, text=text))
                    full_text_parts.append(text)
            return DocumentLoadResult(
                content="\n\n".join(full_text_parts),
                pages=pages,
            )
        except ImportError:
            logger.debug("pdfplumber not available, trying PyPDF2")
        except Exception as e:
            logger.warning("pdfplumber failed: %s, trying PyPDF2", e)

        # Fallback 2: PyPDF2
        try:
            from PyPDF2 import PdfReader  # type: ignore[import-untyped]
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                pages.append(PageInfo(page_number=i, text=text))
                full_text_parts.append(text)
            return DocumentLoadResult(
                content="\n\n".join(full_text_parts),
                pages=pages,
            )
        except ImportError:
            raise ServiceError(
                code=DOCUMENT_PARSE_FAILED,
                message="未安装 PDF 解析库，请安装 PyMuPDF、pdfplumber 或 PyPDF2",
            )

    # ── DOCX ────────────────────────────────────────────────

    def _load_docx(self, file_path: str) -> DocumentLoadResult:
        """使用 python-docx 读取 Word 文档。

        改进点（相比旧版）：
        - 按文档 XML 顺序遍历，不遗漏表格内容
        - 多策略标题检测：outline_level → 样式名(Heading/标题) → 字号+加粗启发式
        - 提取页眉页脚文本
        - 按标题将文档划分为伪页面，改善 chunk 结构感知
        """
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            raise ServiceError(
                code=DOCUMENT_PARSE_FAILED,
                message="未安装 python-docx，请执行: pip install python-docx",
            )

        doc = Document(file_path)

        # 1. 收集页眉页脚
        header_footer_text = self._extract_docx_headers_footers(doc)

        # 2. 按文档顺序遍历 body 子元素（段落 + 表格交替）
        content_blocks: list[str] = []

        if header_footer_text:
            content_blocks.append(header_footer_text)

        content_blocks.extend(self._iter_docx_body_elements(doc))

        # 3. 按标题分节，生成伪页面
        sections = self._split_docx_by_headings(content_blocks)
        pages: list[PageInfo] = []
        for i, section_text in enumerate(sections, start=1):
            clean = section_text.strip()
            if clean:
                pages.append(PageInfo(page_number=i, text=clean))

        content = "\n\n".join(s.strip() for s in sections if s.strip())

        return DocumentLoadResult(
            content=content,
            pages=pages,
        )

    # ── DOCX 辅助方法 ────────────────────────────────────────

    @staticmethod
    def _extract_docx_headers_footers(doc) -> str:
        """提取所有节的页眉和页脚文本。"""
        parts: list[str] = []
        for i, section in enumerate(doc.sections):
            # 第一个 section 总是提取；后续 section 只在未链接到前一节时提取
            is_first = i == 0
            # 页眉
            if section.header:
                if is_first or not section.header.is_linked_to_previous:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
            # 页脚
            if section.footer:
                if is_first or not section.footer.is_linked_to_previous:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)

        if parts:
            return "[文档页眉/页脚]\n" + "\n".join(parts)
        return ""

    @staticmethod
    def _iter_docx_body_elements(doc) -> list[str]:
        """按文档 body XML 子元素顺序遍历，交替返回段落和表格文本。

        关键：通过 doc.element.body 的 child 标签区分 <w:p>（段落）
        和 <w:tbl>（表格），保证表格内容不丢失且位置正确。
        """
        blocks: list[str] = []

        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)

        para_idx = 0
        tbl_idx = 0

        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                if para_idx < len(paragraphs):
                    para = paragraphs[para_idx]
                    para_idx += 1
                    text = para.text.strip()
                    if text:
                        formatted = DocumentLoader._format_docx_paragraph(para, text)
                        blocks.append(formatted)

            elif tag == "tbl":
                if tbl_idx < len(tables):
                    table = tables[tbl_idx]
                    tbl_idx += 1
                    table_text = DocumentLoader._extract_docx_table(table)
                    if table_text:
                        blocks.append(table_text)

        # 防御：如果 XML 遍历遗漏了某些段落（某些 docx 库版本可能不同），回退补充
        while para_idx < len(paragraphs):
            para = paragraphs[para_idx]
            para_idx += 1
            text = para.text.strip()
            if text:
                formatted = DocumentLoader._format_docx_paragraph(para, text)
                blocks.append(formatted)

        return blocks

    @staticmethod
    def _detect_heading_level(para) -> int | None:
        """多策略检测段落标题级别。

        策略优先级：
        1. outline_level 属性（语言无关，最可靠）
        2. 样式名匹配 "Heading N"（英文 Word）
        3. 样式名匹配 "标题 N"（中文 Word）
        4. 样式名包含 "heading"/"title" 关键词
        5. 启发式：字号 >= 16pt 且加粗 → 视为一级标题

        Returns:
            标题级别 1-4，或 None（非标题）。
        """
        # 策略 1: outline_level
        try:
            pf = para.paragraph_format
            if hasattr(pf, "outline_level") and pf.outline_level is not None:
                level = pf.outline_level
                if isinstance(level, int) and 0 <= level <= 8:
                    # outline_level 0 = body text, 1 = Heading 1, ...
                    if level >= 1:
                        return min(level, 4)
        except Exception:
            pass

        # 策略 2 & 3 & 4: 样式名
        style_name = ""
        if para.style and para.style.name:
            style_name = para.style.name.strip()

        if style_name:
            # "Heading 1", "Heading 2", ...
            if style_name.lower().startswith("heading"):
                parts = style_name.split()
                if len(parts) >= 2:
                    try:
                        return min(int(parts[-1]), 4)
                    except ValueError:
                        return 2
                return 2

            # "标题 1", "标题 2", ...
            if "标题" in style_name:
                parts = style_name.split()
                if len(parts) >= 2:
                    try:
                        return min(int(parts[-1]), 4)
                    except ValueError:
                        return 2
                return 2

            # 其他常见标题样式关键词
            style_lower = style_name.lower()
            if any(kw in style_lower for kw in ("title", "head", "subtitle", "h1", "h2", "h3")):
                if "sub" in style_lower:
                    return 3
                return 2

        # 策略 5: 字号 + 加粗启发式
        try:
            font_size = None
            is_bold = False
            for run in para.runs:
                if run.font.size is not None:
                    font_size = run.font.size.pt  # type: ignore[attr-defined]
                if run.bold:
                    is_bold = True
                    break

            if font_size is not None and font_size >= 16 and is_bold:
                return 1
            if font_size is not None and font_size >= 14 and is_bold:
                return 2
        except Exception:
            pass

        return None

    @staticmethod
    def _format_docx_paragraph(para, text: str) -> str:
        """格式化段落文本，为标题添加 markdown 标记。"""
        level = DocumentLoader._detect_heading_level(para)
        if level is not None:
            prefix = "#" * min(level, 4)
            return f"{prefix} {text}"
        return text

    @staticmethod
    def _extract_docx_table(table) -> str:
        """提取 Word 表格内容，格式化为可读文本。

        将表格转为 markdown 风格的文本表示：
        - 第一行视为表头
        - 每行用 " | " 分隔单元格
        """
        rows = table.rows
        if not rows:
            return ""

        lines: list[str] = []
        lines.append("[表格]")

        for row_idx, row in enumerate(rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)

            # 跳过完全为空的行
            if not any(c for c in cells):
                continue

            line = " | ".join(cells)
            lines.append(line)

        if len(lines) <= 1:  # 只有 [表格] 标记，无有效行
            return ""

        return "\n".join(lines)

    @staticmethod
    def _split_docx_by_headings(blocks: list[str]) -> list[str]:
        """按 markdown 标题标记将内容块分节，生成伪页面。

        以 H1/H2（## 及以上）为界切分；如果文档无标题，按块数均分。
        这样每个 section 作为一个 PageInfo，给 chunker 提供更好的结构边界。
        """
        if not blocks:
            return [""]

        sections: list[list[str]] = []
        current: list[str] = []

        for block in blocks:
            is_major_heading = bool(
                (block.startswith("# ") or block.startswith("## "))
                and not block.startswith("### ")
            )
            if is_major_heading and current:
                sections.append(current)
                current = [block]
            else:
                current.append(block)

        if current:
            sections.append(current)

        # 如果分节太少（没有标题），按块数均分，避免单节过大
        if len(sections) <= 1 and len(blocks) > 10:
            sections = []
            section_size = max(5, len(blocks) // 3)
            for i in range(0, len(blocks), section_size):
                sections.append(blocks[i:i + section_size])

        return ["\n\n".join(sec) for sec in sections]

    # ── TXT ─────────────────────────────────────────────────

    def _load_txt(self, file_path: str) -> DocumentLoadResult:
        """读取纯文本文件，自动检测编码。"""
        content = None
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ServiceError(
                code=DOCUMENT_PARSE_FAILED,
                message="无法识别文本编码，请确认文件为 UTF-8 或 GBK 编码",
            )

        return DocumentLoadResult(
            content=content,
            pages=[PageInfo(page_number=1, text=content)],
        )


def compute_file_hash(file_path: str, algorithm: str = "sha256") -> str:
    """计算文件哈希值，用于生成稳定的 document_id。"""
    h = hashlib.new(algorithm)
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()
